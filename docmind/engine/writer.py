"""
Writer v2 — 纯执行器，只执行 fix_plan 中的指令。

核心原则:
- 不做任何判断逻辑（不选值、不决策）
- 不产生任何新值（所有参数来自 fix_plan）
- 只执行 ZIP 级别的 OOXML 修改

用法:
    from pipeline.writer_v2 import apply_fixes

    fix_plan = [
        {"action": "set_sectpr_type", "params": {"section_index": 0, "val": "nextPage"}},
        {"action": "set_header_font", "params": {"header_path": "word/header2.xml", "font_ascii": "Times New Roman", "font_eastAsia": "宋体", "font_size": "18"}},
        {"action": "add_page_number", "params": {"footer_path": "word/footer2.xml", "format": "第{PAGE}页"}},
        {"action": "set_body_font_ascii", "params": {"style_id": "a8", "font_ascii": "Times New Roman"}},
        {"action": "set_style", "params": {"style_id": "1", "font_eastAsia": "黑体", "font_size_pt": 16}},
        {"action": "remove_extra_sectpr", "params": {"section_index": 1}},
    ]
    apply_fixes(fix_plan, Path("output/毕业设计.docx"))
"""

import shutil
import zipfile
import re
from pathlib import Path
from typing import Any, Union

from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"
XML = "http://www.w3.org/XML/1998/namespace"

# ── Action dispatcher ──────────────────────────────────────────

_ACTION_HANDLERS: dict[str, callable] = {}


def _handler(action_name: str):
    """装饰器：注册 action 处理函数"""
    def decorator(fn):
        _ACTION_HANDLERS[action_name] = fn
        return fn
    return decorator


def apply_fixes(fix_plan: list[dict], target_path: Union[str, Path]) -> list[str]:
    """逐条执行 fix_plan 中的每个 action。

    Args:
        fix_plan: action 列表，每个 action 有 "action" 和 "params" 键
        target_path: 目标 docx 文件路径

    Returns:
        执行日志列表，每项描述执行结果
    """
    target_path = Path(target_path)
    logs: list[str] = []

    for i, action in enumerate(fix_plan):
        action_type = action.get("action", "")
        params = action.get("params", {})

        handler = _ACTION_HANDLERS.get(action_type)
        if handler is None:
            logs.append(f"[SKIP #{i}] 未知 action: {action_type}")
            continue

        try:
            handler(target_path, params)
            logs.append(f"[OK #{i}] {action_type} {_brief_params(params)}")
        except Exception as e:
            logs.append(f"[FAIL #{i}] {action_type}: {e}")

    return logs


def _brief_params(params: dict) -> str:
    """生成参数的简要描述"""
    parts = []
    for k, v in params.items():
        if isinstance(v, str) and len(v) > 30:
            v = v[:27] + "..."
        parts.append(f"{k}={v}")
    return " ".join(parts)


# ── ZIP 工具 ──────────────────────────────────────────────────

def _zip_replace(zip_path: Path, internal_path: str, new_data: bytes) -> None:
    """替换 ZIP 包中的单个文件"""
    zip_path_str = str(zip_path)
    tmp = zip_path_str + ".tmp"
    with zipfile.ZipFile(zip_path_str, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == internal_path:
                    data = new_data
                zout.writestr(item, data)
    shutil.move(tmp, zip_path_str)


def _read_zip_entry(zip_path: Path, internal_path: str) -> bytes:
    """读取 ZIP 中的单个条目"""
    with zipfile.ZipFile(zip_path, "r") as z:
        return z.read(internal_path)


def _parse_xml_from_zip(zip_path: Path, internal_path: str):
    """从 ZIP 中读取并解析 XML"""
    return etree.fromstring(_read_zip_entry(zip_path, internal_path))


# ── Action: set_sectpr_type ────────────────────────────────────

@_handler("set_sectpr_type")
def _set_sectpr_type(target_path: Path, params: dict) -> None:
    """修改 sectPr 的 w:type val。

    params:
        section_index: int  — 第几个 sectPr（0-based）
        val: str             — type 值，如 "nextPage", "oddPage", "continuous"
    """
    section_index = params["section_index"]
    val = params["val"]

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    sects = list(root.iter(f"{{{W}}}sectPr"))

    if section_index >= len(sects):
        raise ValueError(f"section_index {section_index} 超出范围 (共 {len(sects)} 个 sectPr)")

    sp = sects[section_index]
    type_el = sp.find(f"{{{W}}}type")
    if type_el is None:
        type_el = etree.SubElement(sp, f"{{{W}}}type")
    type_el.set(f"{{{W}}}val", val)

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: set_page_number_type ─────────────────────────────

@_handler("set_page_number_type")
def _set_page_number_type(target_path: Path, params: dict) -> None:
    """在 sectPr 中设置页码格式（罗马/阿拉伯）和重起序号。

    params:
        section_index: int   — 第几个 sectPr（0-based）
        fmt: str             — "upperRoman" | "decimal"
        start: int | None    — 页码起始值（如 1），None 表示不重起
    """
    section_index = params["section_index"]
    fmt_val = params["fmt"]
    start_val = params.get("start", None)

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    sects = list(root.iter(f"{{{W}}}sectPr"))

    if section_index >= len(sects):
        raise ValueError(f"section_index {section_index} 超出范围 (共 {len(sects)} 个 sectPr)")

    sp = sects[section_index]
    pnt = sp.find(f"{{{W}}}pgNumType")
    if pnt is None:
        pnt = etree.SubElement(sp, f"{{{W}}}pgNumType")
    pnt.set(f"{{{W}}}fmt", fmt_val)
    if start_val is not None:
        pnt.set(f"{{{W}}}start", str(start_val))

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: set_header_font ────────────────────────────────────

@_handler("set_header_font")
def _set_header_font(target_path: Path, params: dict) -> None:
    """修改 header XML 中的 rPr/rFonts。

    params:
        header_path: str     — ZIP 内部路径，如 "word/header2.xml"
        font_ascii: str      — 西文字体名
        font_eastAsia: str   — 中文字体名
        font_size: str       — 字号（half-points），如 "18"
        font_hAnsi: str | None — hAnsi 字体（默认同 font_ascii）
    """
    header_path = params["header_path"]
    font_ascii = params["font_ascii"]
    font_eastAsia = params["font_eastAsia"]
    font_size = params["font_size"]
    font_hAnsi = params.get("font_hAnsi", font_ascii)

    root = _parse_xml_from_zip(target_path, header_path)

    for rPr in root.iter(f"{{{W}}}rPr"):
        old_rf = rPr.find(f"{{{W}}}rFonts")
        if old_rf is not None:
            rPr.remove(old_rf)

        rf = etree.SubElement(rPr, f"{{{W}}}rFonts")
        rf.set(f"{{{W}}}ascii", font_ascii)
        rf.set(f"{{{W}}}hAnsi", font_hAnsi)
        rf.set(f"{{{W}}}eastAsia", font_eastAsia)

        for tag in [f"{{{W}}}sz", f"{{{W}}}szCs"]:
            old_sz = rPr.find(tag)
            if old_sz is not None:
                rPr.remove(old_sz)
            etree.SubElement(rPr, tag, {f"{{{W}}}val": font_size})

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, header_path, new_data)


# ── Action: add_page_number ────────────────────────────────────

@_handler("add_page_number")
def _add_page_number(target_path: Path, params: dict) -> None:
    """在 footer XML 插入 "第{PAGE}页" fldChar 序列。

    params:
        footer_path: str     — ZIP 内部路径，如 "word/footer2.xml"
        format: str          — 页码格式模板，如 "第{PAGE}页"
        font_eastAsia: str | None  — 中文字体
        font_ascii: str | None     — 西文字体
        font_size: str | None      — 字号（half-points），如 "18"
    """
    footer_path = params["footer_path"]
    fmt = params["format"]
    font_eastAsia = params.get("font_eastAsia", "宋体")
    font_ascii = params.get("font_ascii", "Times New Roman")
    font_size = params.get("font_size", "18")

    root = _parse_xml_from_zip(target_path, footer_path)

    for p in list(root):
        root.remove(p)

    p = etree.SubElement(root, f"{{{W}}}p")
    pPr = etree.SubElement(p, f"{{{W}}}pPr")
    alignment = params.get("alignment", "right")  # 中文论文默认右对齐
    etree.SubElement(pPr, f"{{{W}}}jc", {f"{{{W}}}val": alignment})

    parts = fmt.split("{PAGE}")
    for idx, prefix in enumerate(parts):
        if prefix:
            r = etree.SubElement(p, f"{{{W}}}r")
            rPr = etree.SubElement(r, f"{{{W}}}rPr")
            rf = etree.SubElement(rPr, f"{{{W}}}rFonts")
            rf.set(f"{{{W}}}ascii", font_ascii)
            rf.set(f"{{{W}}}eastAsia", font_eastAsia)
            etree.SubElement(rPr, f"{{{W}}}sz", {f"{{{W}}}val": font_size})
            etree.SubElement(rPr, f"{{{W}}}szCs", {f"{{{W}}}val": font_size})
            t = etree.SubElement(r, f"{{{W}}}t", {f"{{{XML}}}space": "preserve"})
            t.text = prefix

        if idx < len(parts) - 1:
            r = etree.SubElement(p, f"{{{W}}}r")
            etree.SubElement(r, f"{{{W}}}fldChar", {f"{{{W}}}fldCharType": "begin"})

            r = etree.SubElement(p, f"{{{W}}}r")
            instr = etree.SubElement(r, f"{{{W}}}instrText", {f"{{{XML}}}space": "preserve"})
            instr.text = " PAGE "

            r = etree.SubElement(p, f"{{{W}}}r")
            etree.SubElement(r, f"{{{W}}}fldChar", {f"{{{W}}}fldCharType": "separate"})

            r = etree.SubElement(p, f"{{{W}}}r")
            rPr = etree.SubElement(r, f"{{{W}}}rPr")
            etree.SubElement(rPr, f"{{{W}}}sz", {f"{{{W}}}val": font_size})
            t = etree.SubElement(r, f"{{{W}}}t")
            t.text = "1"

            r = etree.SubElement(p, f"{{{W}}}r")
            etree.SubElement(r, f"{{{W}}}fldChar", {f"{{{W}}}fldCharType": "end"})

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, footer_path, new_data)


# ── Action: set_body_font_ascii ────────────────────────────────

@_handler("set_body_font_ascii")
def _set_body_font_ascii(target_path: Path, params: dict) -> None:
    """修改 styles.xml 中 body 样式（默认段落样式）的 w:ascii。

    params:
        style_id: str        — 样式 ID，如 "a8"
        font_ascii: str      — 西文字体名
    """
    style_id = params["style_id"]
    font_ascii = params["font_ascii"]

    root = _parse_xml_from_zip(target_path, "word/styles.xml")

    for style in root.iter(f"{{{W}}}style"):
        if style.get(f"{{{W}}}styleId") != style_id:
            continue

        rPr = style.find(f"{{{W}}}rPr")
        if rPr is None:
            rPr = etree.SubElement(style, f"{{{W}}}rPr")

        old_rf = rPr.find(f"{{{W}}}rFonts")
        if old_rf is not None:
            old_rf.set(f"{{{W}}}ascii", font_ascii)
            if not old_rf.get(f"{{{W}}}hAnsi"):
                old_rf.set(f"{{{W}}}hAnsi", font_ascii)
        else:
            rf = etree.SubElement(rPr, f"{{{W}}}rFonts")
            rf.set(f"{{{W}}}ascii", font_ascii)
            rf.set(f"{{{W}}}hAnsi", font_ascii)
        break

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/styles.xml", new_data)


# ── Action: set_style ─────────────────────────────────────────

@_handler("set_style")
def _set_style(target_path: Path, params: dict) -> None:
    """修改 styles.xml 中指定 style_id 的属性。

    params:
        style_id: str            — 样式 ID
        font_eastAsia: str | None
        font_ascii: str | None
        font_hAnsi: str | None
        font_size_pt: float | None  — 字号（point）
        bold: bool | None
        alignment: str | None    — "left", "center", "right", "both"
        before_lines: float | None
        after_lines: float | None
        line: int | None
        lineRule: str | None
    """
    style_id = params["style_id"]

    root = _parse_xml_from_zip(target_path, "word/styles.xml")

    for style in root.iter(f"{{{W}}}style"):
        if style.get(f"{{{W}}}styleId") != style_id:
            continue

        if any(k in params for k in ("font_eastAsia", "font_ascii", "font_hAnsi", "font_size_pt", "bold")):
            rPr = style.find(f"{{{W}}}rPr")
            if rPr is None:
                rPr = etree.SubElement(style, f"{{{W}}}rPr")

            if any(k in params for k in ("font_eastAsia", "font_ascii", "font_hAnsi")):
                old_rf = rPr.find(f"{{{W}}}rFonts")
                if old_rf is not None:
                    rPr.remove(old_rf)
                rf = etree.SubElement(rPr, f"{{{W}}}rFonts")
                if "font_ascii" in params and params["font_ascii"] is not None:
                    rf.set(f"{{{W}}}ascii", params["font_ascii"])
                if "font_hAnsi" in params and params["font_hAnsi"] is not None:
                    rf.set(f"{{{W}}}hAnsi", params["font_hAnsi"])
                elif "font_ascii" in params and params["font_ascii"] is not None:
                    rf.set(f"{{{W}}}hAnsi", params["font_ascii"])
                if "font_eastAsia" in params and params["font_eastAsia"] is not None:
                    rf.set(f"{{{W}}}eastAsia", params["font_eastAsia"])

            if "font_size_pt" in params and params["font_size_pt"] is not None:
                sv = str(int(params["font_size_pt"] * 2))
                for tag in [f"{{{W}}}sz", f"{{{W}}}szCs"]:
                    old = rPr.find(tag)
                    if old is not None:
                        rPr.remove(old)
                    etree.SubElement(rPr, tag, {f"{{{W}}}val": sv})

            if "bold" in params and params["bold"] is not None:
                old_b = rPr.find(f"{{{W}}}b")
                if old_b is not None:
                    rPr.remove(old_b)
                if params["bold"]:
                    etree.SubElement(rPr, f"{{{W}}}b")

        if any(k in params for k in ("alignment", "before_lines", "after_lines", "line", "lineRule")):
            pPr = style.find(f"{{{W}}}pPr")
            if pPr is None:
                pPr = etree.SubElement(style, f"{{{W}}}pPr")

            if "alignment" in params and params["alignment"] is not None:
                old_jc = pPr.find(f"{{{W}}}jc")
                if old_jc is not None:
                    pPr.remove(old_jc)
                etree.SubElement(pPr, f"{{{W}}}jc", {f"{{{W}}}val": params["alignment"]})

            if any(k in params for k in ("before_lines", "after_lines", "line", "lineRule")):
                old_sp = pPr.find(f"{{{W}}}spacing")
                if old_sp is not None:
                    pPr.remove(old_sp)
                sp_attrs = {}
                if "before_lines" in params and params["before_lines"] is not None:
                    sp_attrs[f"{{{W}}}beforeLines"] = str(int(params["before_lines"]))
                if "after_lines" in params and params["after_lines"] is not None:
                    sp_attrs[f"{{{W}}}afterLines"] = str(int(params["after_lines"]))
                if "line" in params and params["line"] is not None:
                    sp_attrs[f"{{{W}}}line"] = str(int(params["line"]))
                if "lineRule" in params and params["lineRule"] is not None:
                    sp_attrs[f"{{{W}}}lineRule"] = params["lineRule"]
                if sp_attrs:
                    etree.SubElement(pPr, f"{{{W}}}spacing", sp_attrs)

        break

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/styles.xml", new_data)


# ── Action: remove_extra_sectpr ────────────────────────────────

@_handler("remove_extra_sectpr")
def _remove_extra_sectpr(target_path: Path, params: dict) -> None:
    """删除嵌入在内容段落中的多余 sectPr。

    params:
        section_index: int | None  — 指定要删除第几个嵌入 sectPr（None=全部删除）
        preserve_body_sectpr: bool — 是否保留 body 末尾的 sectPr（默认 True）
    """
    section_index = params.get("section_index", None)
    preserve_body_sectpr = params.get("preserve_body_sectpr", True)

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    body = root.find(f"{{{W}}}body")
    if body is None:
        return

    removed = 0
    for child in list(body):
        if child.tag != f"{{{W}}}p":
            continue
        pPr = child.find(f"{{{W}}}pPr")
        if pPr is None:
            continue
        extra_sp = pPr.find(f"{{{W}}}sectPr")
        if extra_sp is None:
            continue

        if section_index is not None and removed != section_index:
            removed += 1
            continue

        pPr.remove(extra_sp)
        removed += 1

        if section_index is not None:
            break

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: insert_section_break ───────────────────────────────

@_handler("insert_section_break")
def _insert_section_break(target_path: Path, params: dict) -> None:
    """在指定 H1 标题段落前插入新的 sectPr 分节符段落。

    params:
        heading_text: str  — 需要独立分节符的下一个 H1 标题文字（用于定位）
    """
    heading_text = params["heading_text"]

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    body = root.find(f"{{{W}}}body")

    target_para = None
    target_idx = None
    for i, child in enumerate(list(body)):
        if child.tag == f"{{{W}}}p":
            txt = ''.join(t.text or '' for t in child.iter(f"{{{W}}}t")).strip()
            norm = re.sub(r'\s+', '', txt)
            if norm == re.sub(r'\s+', '', heading_text):
                target_para = child
                target_idx = i
                break

    if target_para is None:
        return

    new_sp = etree.Element(f"{{{W}}}p")
    new_pPr = etree.SubElement(new_sp, f"{{{W}}}pPr")
    sectPr = etree.SubElement(new_pPr, f"{{{W}}}sectPr")
    etree.SubElement(sectPr, f"{{{W}}}type", {f"{{{W}}}val": "nextPage"})
    etree.SubElement(sectPr, f"{{{W}}}pgSz", {f"{{{W}}}w": "11906", f"{{{W}}}h": "16838"})

    target_para.addprevious(new_sp)

    new_data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: create_section_headers ──────────────────────────────

@_handler("create_section_headers")
def _create_section_headers(target_path: Path, params: dict) -> None:
    """为每节创建独立 header 文件，注册到 rels 和 Content_Types，更新 sectPr 引用。

    params:
        section_headers: list of dicts  — 每节的页眉配置
            [{sectpr_index: int,
              default_text: str, even_text: str | None},
             ...]
        font_eastAsia: str   — 中文字体
        font_ascii: str      — 西文字体
        font_size: str       — 字号 half-points
    """
    import zipfile
    from pathlib import Path

    section_headers = params["section_headers"]
    font_eastAsia = params.get("font_eastAsia", "宋体")
    font_ascii = params.get("font_ascii", "Times New Roman")
    font_size = params.get("font_size", "18")

    def _make_hdr_xml(text: str) -> bytes:
        return (
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:hdr xmlns:w="{W}"><w:p><w:pPr><w:jc w:val="center"/>'
            f'<w:pBdr><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
            f'</w:pBdr></w:pPr><w:r><w:rPr>'
            f'<w:rFonts w:eastAsia="{font_eastAsia}" w:ascii="{font_ascii}" w:hAnsi="{font_ascii}"/>'
            f'<w:sz w:val="{font_size}"/><w:szCs w:val="{font_size}"/>'
            f'</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p></w:hdr>'
        ).encode('utf-8')

    tmp_path = str(target_path) + '.hdr_tmp'
    with zipfile.ZipFile(target_path, 'r') as zin:
        doc_xml = etree.parse(zin.open('word/document.xml'))
        rels_xml = etree.parse(zin.open('word/_rels/document.xml.rels'))
        ct_xml = etree.parse(zin.open('[Content_Types].xml'))
        try:
            settings_xml = etree.parse(zin.open('word/settings.xml'))
            has_settings = True
        except KeyError:
            settings_xml = None
            has_settings = False

    rels_root = rels_xml.getroot()
    max_rid = 1
    for rel in rels_root:
        rid = rel.get('Id', '')
        if rid.startswith('rId'):
            try:
                max_rid = max(max_rid, int(rid[3:]))
            except ValueError:
                pass

    existing_headers = set()
    for name in zipfile.ZipFile(target_path, 'r').namelist():
        if 'header' in name and name.endswith('.xml'):
            existing_headers.add(name)

    body = doc_xml.getroot().find(f"{{{W}}}body")
    sects_in_order = []
    for child in body:
        if child.tag == f"{{{W}}}sectPr":
            sects_in_order.append(child)
        elif child.tag == f"{{{W}}}p":
            pp = child.find(f"{{{W}}}pPr")
            if pp is not None:
                sp = pp.find(f"{{{W}}}sectPr")
                if sp is not None:
                    sects_in_order.append(sp)

    ct_root = ct_xml.getroot()
    header_num = 20
    used_texts = {}

    for sh in section_headers:
        si = sh.get("sectpr_index") if "sectpr_index" in sh else sh["section_index"]
        if si >= len(sects_in_order):
            continue
        sp = sects_in_order[si]

        default_text = sh.get("default_text", "")
        even_text = sh.get("even_text")

        if not default_text:
            continue

        key = (default_text, even_text or "")
        if key not in used_texts:
            max_rid += 1
            header_num += 1
            def_fname = f'header{header_num}.xml'
            def_rid = f'rId{max_rid}'

            ct_root.append(ct_root.makeelement(f'{{{CT}}}Override', {
                'PartName': f'/word/{def_fname}',
                'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml'
            }))
            rels_root.append(rels_root.makeelement('Relationship', {
                'Id': def_rid,
                'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header',
                'Target': def_fname,
                'TargetMode': 'Internal'
            }))

            eve_fname = eve_rid = None
            if even_text:
                max_rid += 1
                header_num += 1
                eve_fname = f'header{header_num}.xml'
                eve_rid = f'rId{max_rid}'
                ct_root.append(ct_root.makeelement(f'{{{CT}}}Override', {
                    'PartName': f'/word/{eve_fname}',
                    'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml'
                }))
                rels_root.append(rels_root.makeelement('Relationship', {
                    'Id': eve_rid,
                    'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header',
                    'Target': eve_fname,
                    'TargetMode': 'Internal'
                }))

            used_texts[key] = (def_fname, def_rid, eve_fname, eve_rid)

        def_fname, def_rid, eve_fname, eve_rid = used_texts[key]

        # 清除旧引用 + 旧标志
        for old in list(sp.findall(f"{{{W}}}headerReference")):
            sp.remove(old)
        for old_flag in list(sp.findall(f"{{{W}}}evenAndOddHeaders")):
            sp.remove(old_flag)
        for old_tp in list(sp.findall(f"{{{W}}}titlePg")):
            sp.remove(old_tp)

        # ── 清理：如果 sectPr 在内容段落中，移到独立空段 ──
        parent = sp.getparent()  # pPr
        if parent is not None:
            grandparent = parent.getparent()  # p (the paragraph)
            if grandparent is not None and grandparent.tag == f'{{{W}}}p':
                para_txt = ''.join(t.text or '' for t in grandparent.iter(f'{{{W}}}t')).strip()
                if para_txt:
                    parent.remove(sp)
                    new_para = etree.Element(f'{{{W}}}p')
                    new_pPr = etree.SubElement(new_para, f'{{{W}}}pPr')
                    new_pPr.append(sp)
                    grandparent.addnext(new_para)
                    sp = new_pPr.find(f'{{{W}}}sectPr')  # 后续操作使用移动后的 sectPr

        hr = etree.SubElement(sp, f"{{{W}}}headerReference")
        hr.set(f"{{{W}}}type", 'default')
        hr.set(f"{{{R}}}id", def_rid)

        if eve_rid:
            hr = etree.SubElement(sp, f"{{{W}}}headerReference")
            hr.set(f"{{{W}}}type", 'even')
            hr.set(f"{{{R}}}id", eve_rid)
            if sp.find(f"{{{W}}}evenAndOddHeaders") is None:
                etree.SubElement(sp, f"{{{W}}}evenAndOddHeaders")

    # ── 清理文档级 settings.xml 中的全局 evenAndOddHeaders ──
    settings_data = None
    if has_settings and settings_xml is not None:
        settings_root = settings_xml.getroot()
        for eoh in settings_root.findall(f'{{{W}}}evenAndOddHeaders'):
            settings_root.remove(eoh)
        settings_data = etree.tostring(settings_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    doc_data = etree.tostring(doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
    rels_data = etree.tostring(rels_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
    ct_data = etree.tostring(ct_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    with zipfile.ZipFile(target_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = doc_data
                elif item.filename == 'word/_rels/document.xml.rels':
                    data = rels_data
                elif item.filename == '[Content_Types].xml':
                    data = ct_data
                elif item.filename == 'word/settings.xml' and settings_data is not None:
                    data = settings_data
                zout.writestr(item, data)
            for (dt, et), (df, dr, ef, er) in used_texts.items():
                zout.writestr(zipfile.ZipInfo(f'word/{df}'),
                              _make_hdr_xml_via_python_docx(dt, font_size, font_eastAsia, font_ascii))
                if ef:
                    zout.writestr(zipfile.ZipInfo(f'word/{ef}'),
                                  _make_hdr_xml_via_python_docx(et, font_size, font_eastAsia, font_ascii))

    shutil.move(tmp_path, str(target_path))


@_handler("set_header_text")
def _set_header_text(target_path: Path, params: dict) -> None:
    """替换 header XML 中的所有文字内容，保留字体格式。

    params:
        header_path: str     — ZIP 内部路径，如 "word/header2.xml"
        text: str            — 新的页眉文字
        font_eastAsia: str   — 中文字体（默认 "宋体"）
        font_ascii: str      — 西文字体（默认 "Times New Roman"）
        font_size: str       — 字号 half-points（默认 "18"）
    """
    header_path = params["header_path"]
    text = params["text"]
    font_eastAsia = params.get("font_eastAsia", "宋体")
    font_ascii = params.get("font_ascii", "Times New Roman")
    font_size = params.get("font_size", "18")

    root = _parse_xml_from_zip(target_path, header_path)

    for p in list(root):
        root.remove(p)

    p = etree.SubElement(root, f"{{{W}}}p")
    pPr = etree.SubElement(p, f"{{{W}}}pPr")
    etree.SubElement(pPr, f"{{{W}}}jc", {f"{{{W}}}val": "center"})
    etree.SubElement(pPr, f"{{{W}}}pBdr").append(
        etree.Element(f"{{{W}}}bottom", {
            f"{{{W}}}val": "single",
            f"{{{W}}}sz": "6",
            f"{{{W}}}space": "1",
            f"{{{W}}}color": "auto",
        })
    )

    r = etree.SubElement(p, f"{{{W}}}r")
    rPr = etree.SubElement(r, f"{{{W}}}rPr")
    rf = etree.SubElement(rPr, f"{{{W}}}rFonts")
    rf.set(f"{{{W}}}eastAsia", font_eastAsia)
    rf.set(f"{{{W}}}ascii", font_ascii)
    rf.set(f"{{{W}}}hAnsi", font_ascii)
    etree.SubElement(rPr, f"{{{W}}}sz", {f"{{{W}}}val": font_size})
    etree.SubElement(rPr, f"{{{W}}}szCs", {f"{{{W}}}val": font_size})
    t = etree.SubElement(r, f"{{{W}}}t", {f"{{{XML}}}space": "preserve"})
    t.text = text

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, header_path, new_data)



# ── Action: clear_section_headers ────────────────────────────

@_handler("clear_section_headers")
def _clear_section_headers(target_path: Path, params: dict) -> None:
    """移除 sectPr 中所有 headerReference。"""
    section_index = params["section_index"]

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    sects = list(root.iter(f"{{{W}}}sectPr"))
    if section_index >= len(sects):
        return

    sp = sects[section_index]
    for hr in list(sp.findall(f"{{{W}}}headerReference")):
        sp.remove(hr)
    # 同时移除 evenAndOddHeaders
    eoh = sp.find(f"{{{W}}}evenAndOddHeaders")
    if eoh is not None:
        sp.remove(eoh)

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: add_footer_reference ─────────────────────────────

@_handler("add_footer_reference")
def _add_footer_reference(target_path: Path, params: dict) -> None:
    """在 sectPr 中添加默认 footerReference。"""
    import zipfile
    section_index = params["section_index"]

    rels_root = _parse_xml_from_zip(target_path, "word/_rels/document.xml.rels")
    footer_rid = None
    for rel in rels_root:
        if "footer" in (rel.get("Type", "")):
            footer_rid = rel.get("Id")
            break
    if footer_rid is None:
        return

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    sects = list(root.iter(f"{{{W}}}sectPr"))
    if section_index >= len(sects):
        return

    sp = sects[section_index]
    title_pg = sp.find(f"{{{W}}}titlePg")
    if title_pg is not None:
        sp.remove(title_pg)

    if sp.findall(f"{{{W}}}footerReference"):
        return

    fr = etree.SubElement(sp, f"{{{W}}}footerReference")
    fr.set(f"{{{W}}}type", "default")
    fr.set(f"{{{R}}}id", footer_rid)

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: set_code_block_style ──────────────────────────────

@_handler("set_code_block_style")
def _set_code_block_style(target_path: Path, params: dict) -> None:
    """为指定的段落设置代码块样式：灰色背景 + 等宽字体。

    params:
        paragraph_indices: list[int]  — 段落索引列表（0-based，body 子元素中的 p 元素序号）
    """
    paragraph_indices = params["paragraph_indices"]

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    body = root.find(f"{{{W}}}body")
    if body is None:
        return

    # 收集所有 p 段落的索引映射
    p_elements = []
    for i, child in enumerate(body):
        if child.tag == f"{{{W}}}p":
            p_elements.append((i, child))

    for pi in paragraph_indices:
        if pi >= len(p_elements):
            continue
        _, p = p_elements[pi]
        # 设置段落背景色
        pPr = p.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(p, f"{{{W}}}pPr")
            p.insert(0, pPr)
        # 移除旧 shd
        for old_shd in pPr.findall(f"{{{W}}}shd"):
            pPr.remove(old_shd)
        etree.SubElement(pPr, f"{{{W}}}shd", {
            f"{{{W}}}val": "clear",
            f"{{{W}}}fill": "F5F5F5",
        })

        # 设置每个 run 的等宽字体
        for run in p.iter(f"{{{W}}}r"):
            rPr = run.find(f"{{{W}}}rPr")
            if rPr is None:
                rPr = etree.SubElement(run, f"{{{W}}}rPr")
                run.insert(0, rPr)
            old_rf = rPr.find(f"{{{W}}}rFonts")
            if old_rf is not None:
                old_rf.set(f"{{{W}}}ascii", "Consolas")
                old_rf.set(f"{{{W}}}hAnsi", "Consolas")
            else:
                rf = etree.SubElement(rPr, f"{{{W}}}rFonts")
                rf.set(f"{{{W}}}ascii", "Consolas")
                rf.set(f"{{{W}}}hAnsi", "Consolas")
                rPr.insert(0, rf)

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: insert_toc ────────────────────────────────────────

@_handler("insert_toc")
def _insert_toc(target_path: Path, params: dict) -> None:
    """在指定位置插入 TOC（目录）域。

    params:
        after_paragraph_index: int | None  — 插入到该段落后（None = prepend 到 body 开头）
    """
    after_paragraph_index = params.get("after_paragraph_index", None)

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    body = root.find(f"{{{W}}}body")
    if body is None:
        return

    # 收集所有 p 段落的索引映射
    p_elements = []
    for i, child in enumerate(body):
        if child.tag == f"{{{W}}}p":
            p_elements.append((i, child))

    # 构建 TOC 域段落
    toc_p = etree.SubElement(body, f"{{{W}}}p")  # 临时添加
    # 先移除，稍后插入到正确位置
    body.remove(toc_p)

    # pPr: 设置段落格式
    pPr = etree.SubElement(toc_p, f"{{{W}}}pPr")
    etree.SubElement(pPr, f"{{{W}}}spacing", {
        f"{{{W}}}before": "120",
        f"{{{W}}}after": "120",
    })

    # TOC begin
    r_begin = etree.SubElement(toc_p, f"{{{W}}}r")
    etree.SubElement(r_begin, f"{{{W}}}fldChar", {f"{{{W}}}fldCharType": "begin"})

    # TOC instrText
    r_instr = etree.SubElement(toc_p, f"{{{W}}}r")
    instr = etree.SubElement(r_instr, f"{{{W}}}instrText", {f"{{{XML}}}space": "preserve"})
    instr.text = " TOC \\o \"1-3\" \\h \\z \\u "

    # TOC separate
    r_sep = etree.SubElement(toc_p, f"{{{W}}}r")
    etree.SubElement(r_sep, f"{{{W}}}fldChar", {f"{{{W}}}fldCharType": "separate"})

    # TOC placeholder text
    r_text = etree.SubElement(toc_p, f"{{{W}}}r")
    t = etree.SubElement(r_text, f"{{{W}}}t")
    t.text = '（请右键点击此处，选择\u201c更新域\u201d以生成目录）'

    # TOC end
    r_end = etree.SubElement(toc_p, f"{{{W}}}r")
    etree.SubElement(r_end, f"{{{W}}}fldChar", {f"{{{W}}}fldCharType": "end"})

    # 在指定位置插入 TOC 段落
    if after_paragraph_index is None:
        # 插入到 body 最前面
        body.insert(0, toc_p)
    else:
        if after_paragraph_index < len(p_elements):
            _, target_para = p_elements[after_paragraph_index]
            target_para.addnext(toc_p)
        else:
            body.insert(0, toc_p)

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── Action: insert_red_header ──────────────────────────────────

@_handler("insert_red_header")
def _insert_red_header(target_path: Path, params: dict) -> None:
    """在文档正文开头插入公文红头区域。

    依据 GB/T 9704 格式：
    - 第一段：发文机关全称（红色大号居中 + 下方红色反线）
    - 第二段：发文字号（红色小号居中）

    params:
        org_name: str       — 发文机关全称
        doc_number: str     — 发文字号
        font_eastAsia: str  — 中文字体（默认 "方正小标宋简体"）
        org_font_size: str  — 机关名称字号 half-points（默认 "84" = 42pt 初号）
        doc_font_size: str  — 发文字号字号 half-points（默认 "32" = 16pt 三号）
    """
    org_name = params["org_name"]
    doc_number = params["doc_number"]
    font_east = params.get("font_eastAsia", "方正小标宋简体")
    org_font_size = params.get("org_font_size", "84")
    doc_font_size = params.get("doc_font_size", "32")

    root = _parse_xml_from_zip(target_path, "word/document.xml")
    body = root.find(f"{{{W}}}body")
    if body is None:
        return

    # ── 机关名称段落（带红色反线 = 下边框） ──
    p_org = etree.Element(f"{{{W}}}p")
    pPr_org = etree.SubElement(p_org, f"{{{W}}}pPr")
    etree.SubElement(pPr_org, f"{{{W}}}jc", {f"{{{W}}}val": "center"})
    # 段前间距，让红头离页面上边距近一些
    etree.SubElement(pPr_org, f"{{{W}}}spacing", {
        f"{{{W}}}before": "0",
        f"{{{W}}}after": "0",
    })
    # 红色反线（段落下边框）
    pBdr = etree.SubElement(pPr_org, f"{{{W}}}pBdr")
    etree.SubElement(pBdr, f"{{{W}}}bottom", {
        f"{{{W}}}val": "single",
        f"{{{W}}}sz": "12",
        f"{{{W}}}space": "4",
        f"{{{W}}}color": "FF0000",
    })
    r_org = etree.SubElement(p_org, f"{{{W}}}r")
    rPr_org = etree.SubElement(r_org, f"{{{W}}}rPr")
    rf_org = etree.SubElement(rPr_org, f"{{{W}}}rFonts")
    rf_org.set(f"{{{W}}}eastAsia", font_east)
    etree.SubElement(rPr_org, f"{{{W}}}sz", {f"{{{W}}}val": org_font_size})
    etree.SubElement(rPr_org, f"{{{W}}}szCs", {f"{{{W}}}val": org_font_size})
    etree.SubElement(rPr_org, f"{{{W}}}color", {f"{{{W}}}val": "FF0000"})
    t_org = etree.SubElement(r_org, f"{{{W}}}t", {f"{{{XML}}}space": "preserve"})
    t_org.text = org_name

    # ── 发文字号段落 ──
    p_doc = etree.Element(f"{{{W}}}p")
    pPr_doc = etree.SubElement(p_doc, f"{{{W}}}pPr")
    etree.SubElement(pPr_doc, f"{{{W}}}jc", {f"{{{W}}}val": "center"})
    etree.SubElement(pPr_doc, f"{{{W}}}spacing", {
        f"{{{W}}}before": "60",
        f"{{{W}}}after": "0",
    })
    r_doc = etree.SubElement(p_doc, f"{{{W}}}r")
    rPr_doc = etree.SubElement(r_doc, f"{{{W}}}rPr")
    rf_doc = etree.SubElement(rPr_doc, f"{{{W}}}rFonts")
    rf_doc.set(f"{{{W}}}eastAsia", font_east)
    etree.SubElement(rPr_doc, f"{{{W}}}sz", {f"{{{W}}}val": doc_font_size})
    etree.SubElement(rPr_doc, f"{{{W}}}szCs", {f"{{{W}}}val": doc_font_size})
    etree.SubElement(rPr_doc, f"{{{W}}}color", {f"{{{W}}}val": "FF0000"})
    t_doc = etree.SubElement(r_doc, f"{{{W}}}t", {f"{{{XML}}}space": "preserve"})
    t_doc.text = doc_number

    # 插入到 body 开头
    if len(body) > 0:
        body.insert(0, p_doc)
        body.insert(0, p_org)
    else:
        body.append(p_org)
        body.append(p_doc)

    new_data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _zip_replace(target_path, "word/document.xml", new_data)


# ── 便捷函数 ──────────────────────────────────────────────────

def fix_from_reconciler_output(reconciler_output: dict, target_path: Union[str, Path]) -> list[str]:
    raise NotImplementedError("请在 agent.py 中使用 _reconciler_to_fix_plan()")


def _make_hdr_xml_via_python_docx(text: str, font_size: str = "21",
                                  font_eastAsia: str = "宋体",
                                  font_ascii: str = "Times New Roman") -> bytes:
    """用 python-docx 创建临时文档，提取其 header XML。字号/字体来自规范，不硬编码。"""
    try:
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile, os

        doc = Document()
        sec = doc.sections[0]
        sec.header.is_linked_to_previous = False
        hdr = sec.header
        p = hdr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run(text)
        rPr = run._r.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_eastAsia}" w:ascii="{font_ascii}" w:hAnsi="{font_ascii}"/>')
        rPr.insert(0, rFonts)
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{font_size}"/>')
        rPr.append(sz)
        szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{font_size}"/>')
        rPr.append(szCs)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
        pPr.append(pBdr)

        fd, tmp = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        doc.save(tmp)
        with zipfile.ZipFile(tmp, 'r') as z:
            for n in z.namelist():
                if n.startswith('word/header') and n.endswith('.xml'):
                    hdr_xml = z.read(n)
                    break
            else:
                hdr_xml = _make_hdr_xml(text)
        os.unlink(tmp)
        return hdr_xml
    except Exception:
        return _make_hdr_xml(text)
