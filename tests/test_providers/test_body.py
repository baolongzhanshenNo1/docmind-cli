"""Tests for BodyProvider."""
from generator.providers.body import BodyProvider
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TestBodyProvider:
    """Unit tests for BodyProvider.render()."""

    def test_section_type_is_body(self):
        assert BodyProvider.section_type == "body"

    def test_render_produces_heading_only_when_no_content(self, blank_doc):
        provider = BodyProvider()
        section = {"title": "绪论"}
        config = {}

        provider.render(blank_doc, section, config)

        # Only 1 heading, no placeholder paragraphs
        assert len(blank_doc.paragraphs) == 1

    def test_heading_is_centered_and_bold(self, blank_doc):
        provider = BodyProvider()
        section = {"title": "文献综述"}
        config = {}

        provider.render(blank_doc, section, config)

        heading_para = blank_doc.paragraphs[0]
        assert heading_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert len(heading_para.runs) == 1
        assert heading_para.runs[0].text == "文献综述"
        assert heading_para.runs[0].bold is True

    def test_heading_font_size_defaults_to_16pt(self, blank_doc):
        provider = BodyProvider()
        section = {"title": "实验设计"}
        config = {}

        provider.render(blank_doc, section, config)
        heading_para = blank_doc.paragraphs[0]
        from docx.shared import Pt
        assert heading_para.runs[0].font.size == Pt(16)

    def test_heading_uses_config_font_size(self, blank_doc):
        provider = BodyProvider()
        section = {"title": "方法"}
        config = {"fonts": {"heading": {"size": 20}}}

        provider.render(blank_doc, section, config)
        heading_para = blank_doc.paragraphs[0]
        from docx.shared import Pt
        assert heading_para.runs[0].font.size == Pt(20)

    def test_body_paragraphs_render_content(self, blank_doc):
        provider = BodyProvider()
        section = {"title": "数据分析", "content": "第一段内容。\n\n第二段内容。\n\n第三段内容。"}
        config = {}

        provider.render(blank_doc, section, config)

        # 1 heading + 3 content paragraphs = 4 total
        assert len(blank_doc.paragraphs) == 4
        assert "第一段内容" in blank_doc.paragraphs[1].text
        assert "第二段内容" in blank_doc.paragraphs[2].text
        assert "第三段内容" in blank_doc.paragraphs[3].text

    def test_body_paragraphs_not_empty_when_content_set(self, blank_doc):
        provider = BodyProvider()
        section = {"title": "结论与展望", "content": "本章总结了研究成果。"}
        config = {}

        provider.render(blank_doc, section, config)

        # Heading + 1 content paragraph
        assert len(blank_doc.paragraphs) == 2
        assert len(blank_doc.paragraphs[1].text) > 0
