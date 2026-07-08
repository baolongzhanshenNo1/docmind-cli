"""Tests for BodyProvider subsections / figures / tables rendering."""
from docx.shared import Pt
from generator.providers.body import BodyProvider


class TestBodySubsections:
    """Verify that BodyProvider recursively renders subsections."""

    def test_subsection_title_appears_in_document(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "主章节",
            "content": "",
            "subsections": [
                {"title": "子节 1.1", "content": ""},
                {"title": "子节 1.2", "content": ""},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        # Main heading + sub-heading 1 + sub-heading 2 = 3 paragraphs
        texts = [p.text for p in blank_doc.paragraphs]
        assert "主章节" in texts
        assert "子节 1.1" in texts
        assert "子节 1.2" in texts

    def test_subsection_heading_is_14pt_bold(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "主章节",
            "subsections": [
                {"title": "子节 A", "content": ""},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        # Paragraph 1 = main heading (16pt), paragraph 2 = sub heading (14pt)
        sub_para = blank_doc.paragraphs[1]
        assert sub_para.runs[0].text == "子节 A"
        assert sub_para.runs[0].font.size == Pt(14)
        assert sub_para.runs[0].bold is True

    def test_subsection_has_indent(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "主章节",
            "subsections": [
                {"title": "缩进子节", "content": "子节正文内容。"},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        # Sub-heading: left indent should be set for level >= 2
        sub_para = blank_doc.paragraphs[1]
        assert sub_para.paragraph_format.left_indent == Pt(20)

    def test_deeply_nested_subsections(self, blank_doc):
        """Subsections can be arbitrarily nested."""
        provider = BodyProvider()
        section = {
            "title": "L1",
            "subsections": [
                {
                    "title": "L2",
                    "content": "",
                    "subsections": [
                        {"title": "L3", "content": "L3 内容。"},
                    ],
                },
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        texts = [p.text for p in blank_doc.paragraphs]
        assert "L1" in texts
        assert "L2" in texts
        assert "L3" in texts
        assert "L3 内容。" in texts

    def test_subsection_content_gets_indented(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "主章节",
            "subsections": [
                {"title": "子节", "content": "子节的第一段。\n\n子节的第二段。"},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        # Sub-heading should be para 1, content paras 2 and 3
        content_para = blank_doc.paragraphs[2]
        assert "子节的第一段" in content_para.text
        assert content_para.paragraph_format.left_indent == Pt(20)


class TestBodyFigures:
    """Verify figure list rendering."""

    def test_figures_rendered_as_labeled_list(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "实验结果",
            "figures": [
                {"title": "系统架构图", "caption": "系统整体架构示意"},
                {"title": "流程图", "caption": "数据处理流程"},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        texts = [p.text for p in blank_doc.paragraphs]
        assert "图 1 系统架构图: 系统整体架构示意" in texts
        assert "图 2 流程图: 数据处理流程" in texts

    def test_figure_without_title(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "测试",
            "figures": [
                {"caption": "仅标题的图"},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        texts = [p.text for p in blank_doc.paragraphs]
        assert "图 1: 仅标题的图" in texts


class TestBodyTables:
    """Verify table rendering."""

    def test_table_with_headers_and_rows(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "数据汇总",
            "tables": [
                {
                    "title": "性能对比表",
                    "headers": ["方法", "准确率", "耗时(ms)"],
                    "rows": [
                        ["方法A", "95.2%", "120"],
                        ["方法B", "97.8%", "85"],
                    ],
                    "caption": "各方法在测试集上的表现",
                },
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        # Should have: heading, table title, table (not a paragraph), caption
        texts = [p.text for p in blank_doc.paragraphs]
        assert "数据汇总" in texts
        assert "性能对比表" in texts
        assert "表 1: 各方法在测试集上的表现" in texts

        # Verify the actual table content
        tables = blank_doc.tables
        assert len(tables) == 1
        table = tables[0]
        assert len(table.rows) == 3  # 1 header + 2 data rows
        assert table.rows[0].cells[0].text == "方法"
        assert table.rows[0].cells[1].text == "准确率"
        assert table.rows[0].cells[2].text == "耗时(ms)"
        assert table.rows[1].cells[0].text == "方法A"
        assert table.rows[2].cells[1].text == "97.8%"

    def test_table_without_caption(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "简要数据",
            "tables": [
                {
                    "headers": ["Key", "Value"],
                    "rows": [["foo", "bar"]],
                },
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        # Only one table, no caption paragraph
        assert len(blank_doc.tables) == 1


class TestBodyFullIntegration:
    """End-to-end: content + figures + tables + subsections together."""

    def test_full_section_with_everything(self, blank_doc):
        provider = BodyProvider()
        section = {
            "title": "综合章节",
            "content": "这是章节引言。\n\n这是第二段引言。",
            "figures": [
                {"title": "架构图", "caption": "系统架构"},
            ],
            "tables": [
                {
                    "title": "对比表",
                    "headers": ["A", "B"],
                    "rows": [["1", "2"]],
                    "caption": "对比结果",
                },
            ],
            "subsections": [
                {"title": "子节一", "content": "详细讨论。"},
            ],
        }
        config = {}

        provider.render(blank_doc, section, config)

        texts = [p.text for p in blank_doc.paragraphs]

        # Heading
        assert "综合章节" in texts
        # Content paragraphs
        assert any("章节引言" in t for t in texts)
        assert any("第二段引言" in t for t in texts)
        # Figure
        assert any("图 1" in t for t in texts)
        # Table title + caption
        assert any("对比表" in t for t in texts)
        assert any("表 1: 对比结果" in t for t in texts)
        # Subsection
        assert "子节一" in texts
        assert "详细讨论。" in texts

        # Verify table was created
        assert len(blank_doc.tables) == 1
