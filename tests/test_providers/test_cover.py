"""Tests for CoverProvider."""
from generator.providers.cover import CoverProvider
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TestCoverProvider:
    """Unit tests for CoverProvider.render()."""

    def test_section_type_is_cover(self):
        assert CoverProvider.section_type == "cover"

    def test_render_produces_paragraphs(self, blank_doc):
        provider = CoverProvider()
        section = {"title": "测试论文标题"}
        config = {}

        # Before render: Document has 0 paragraphs (python-docx blank doc)
        assert len(blank_doc.paragraphs) == 0

        provider.render(blank_doc, section, config)

        # 4 empty + 1 title + 4 empty + 1 footer = 10 paragraphs
        assert len(blank_doc.paragraphs) == 10

    def test_title_paragraph_is_centered_and_bold(self, blank_doc):
        provider = CoverProvider()
        section = {"title": "人工智能研究"}
        config = {}

        provider.render(blank_doc, section, config)

        # Title is paragraph at index 4 (0-3 are empty paragraphs)
        title_para = blank_doc.paragraphs[4]
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert len(title_para.runs) == 1
        assert title_para.runs[0].text == "人工智能研究"
        assert title_para.runs[0].bold is True

    def test_title_font_size_is_26pt(self, blank_doc):
        provider = CoverProvider()
        section = {"title": "测试"}
        config = {}

        provider.render(blank_doc, section, config)
        title_para = blank_doc.paragraphs[4]
        from docx.shared import Pt
        assert title_para.runs[0].font.size == Pt(26)

    def test_first_four_paragraphs_are_empty(self, blank_doc):
        provider = CoverProvider()
        section = {"title": "任意标题"}
        config = {}

        provider.render(blank_doc, section, config)
        for i in range(4):
            assert blank_doc.paragraphs[i].text == ""

    def test_footer_paragraph_content(self, blank_doc):
        provider = CoverProvider()
        section = {"title": "标题"}
        config = {}

        provider.render(blank_doc, section, config)
        # Last paragraph is at index 9
        footer_para = blank_doc.paragraphs[9]
        assert "（封面 — 无页码）" in footer_para.text
        assert footer_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
