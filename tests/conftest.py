"""pytest fixtures shared across DocMind 2.0 unit tests."""
import pytest
from docx import Document


@pytest.fixture
def blank_doc():
    """Return a fresh, blank python-docx Document."""
    return Document()


@pytest.fixture
def doc_with_sections():
    """Return a Document with 3 distinct sections (breaks)."""
    doc = Document()
    # Section 1: default
    doc.add_paragraph("Section 1 content")
    doc.add_section()
    # Section 2
    doc.add_paragraph("Section 2 content")
    doc.add_section()
    # Section 3
    doc.add_paragraph("Section 3 content")
    return doc
