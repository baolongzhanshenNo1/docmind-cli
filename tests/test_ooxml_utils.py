"""Unit tests for ooxml_utils.py — OOXML pure functions."""
import pytest
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from generator.ooxml_utils import (
    set_pgNumType,
    remove_pgNumType,
    set_page_start,
    remove_page_start,
    fix_section_breaks,
)


class TestPgNumType:
    """set_pgNumType / remove_pgNumType."""

    def test_set_pgNumType_creates_element(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        # Initially no pgNumType
        sectPr = section._sectPr
        assert sectPr.find(qn("w:pgNumType")) is None

        set_pgNumType(section, "arabic")
        pg = sectPr.find(qn("w:pgNumType"))
        assert pg is not None
        assert pg.get(qn("w:fmt")) == "decimal"

    def test_set_pgNumType_updates_existing(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        set_pgNumType(section, "arabic")
        set_pgNumType(section, "roman_upper")
        pg = section._sectPr.find(qn("w:pgNumType"))
        assert pg.get(qn("w:fmt")) == "upperRoman"

    def test_set_pgNumType_roman_lower(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        set_pgNumType(section, "roman_lower")
        pg = section._sectPr.find(qn("w:pgNumType"))
        assert pg.get(qn("w:fmt")) == "lowerRoman"

    def test_set_pgNumType_unknown_fmt_defaults_to_decimal(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        set_pgNumType(section, "unknown_fmt")
        pg = section._sectPr.find(qn("w:pgNumType"))
        assert pg.get(qn("w:fmt")) == "decimal"

    def test_remove_pgNumType_removes(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        set_pgNumType(section, "arabic")
        assert section._sectPr.find(qn("w:pgNumType")) is not None
        remove_pgNumType(section)
        assert section._sectPr.find(qn("w:pgNumType")) is None

    def test_remove_pgNumType_noop_when_missing(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        assert section._sectPr.find(qn("w:pgNumType")) is None
        # Should not raise
        remove_pgNumType(section)
        assert section._sectPr.find(qn("w:pgNumType")) is None


class TestPageStart:
    """set_page_start / remove_page_start."""

    def test_set_page_start_creates_and_sets(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        assert section._sectPr.find(qn("w:pgNumType")) is None

        set_page_start(section, 5)
        pg = section._sectPr.find(qn("w:pgNumType"))
        assert pg is not None
        assert pg.get(qn("w:start")) == "5"

    def test_set_page_start_overwrites_existing(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        set_page_start(section, 1)
        set_page_start(section, 10)
        pg = section._sectPr.find(qn("w:pgNumType"))
        assert pg.get(qn("w:start")) == "10"
        # fmt should not be touched if pgNumType already existed
        assert pg.get(qn("w:fmt")) is None

    def test_remove_page_start_removes_attribute(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        set_page_start(section, 7)
        assert section._sectPr.find(qn("w:pgNumType")).get(qn("w:start")) == "7"
        remove_page_start(section)
        pg = section._sectPr.find(qn("w:pgNumType"))
        # pgNumType element should still exist, just without start attr
        assert pg is not None
        assert pg.get(qn("w:start")) is None

    def test_remove_page_start_noop_when_no_start(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        # Create pgNumType without start attr
        pg = OxmlElement("w:pgNumType")
        section._sectPr.insert(0, pg)
        # Should not raise
        remove_page_start(section)
        assert section._sectPr.find(qn("w:pgNumType")) is not None

    def test_remove_page_start_noop_when_no_pgNumType(self, doc_with_sections):
        section = doc_with_sections.sections[0]
        assert section._sectPr.find(qn("w:pgNumType")) is None
        # Should not raise
        remove_page_start(section)


class TestFixSectionBreaks:
    """fix_section_breaks post-processing."""

    def test_fix_section_breaks_noop_on_simple_doc(self, blank_doc):
        """A simple doc with no sectPr in paragraphs should not raise."""
        blank_doc.add_paragraph("Hello")
        blank_doc.add_paragraph("World")
        fix_section_breaks(blank_doc)
        # Just verify no exception
        assert len(blank_doc.paragraphs) == 2

    def test_fix_section_breaks_reorders_children(self, blank_doc):
        """Verify sectPr children are reordered per SECTPR_ORDER."""
        blank_doc.add_paragraph("Content")
        # Access the single default section's sectPr and add children
        # in a deliberately wrong order.
        section = blank_doc.sections[0]
        sectPr = section._sectPr

        # Add elements in wrong order: cols before pgSz
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), "2")
        pgSz = OxmlElement("w:pgSz")
        pgSz.set(qn("w:w"), "11906")
        sectPr.append(cols)
        sectPr.append(pgSz)

        fix_section_breaks(blank_doc)

        # After fix, pgSz should come before cols (based on SECTPR_ORDER)
        children = list(sectPr)
        tags = [c.tag.split("}")[-1] for c in children]
        pgSz_idx = tags.index("pgSz")
        cols_idx = tags.index("cols")
        assert pgSz_idx < cols_idx, f"Expected pgSz before cols, got {tags}"

    def test_fix_section_breaks_moves_sectpr_from_empty_para(self):
        """sectPr on an empty paragraph should be moved to the previous
        content paragraph (which also has sectPr), and the empty
        paragraph removed.  This mirrors what python-docx produces
        when add_section() is called: the last paragraph of the old
        section carries a sectPr, and the new section's first
        (empty) paragraph gets its own sectPr."""
        from docx import Document

        doc = Document()
        # Add a content paragraph that has its own sectPr
        doc.add_paragraph("Chapter 1")
        body = doc.element.body
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        # Give the content paragraph a sectPr
        content_p = body.find(f"{{{W}}}p")
        content_pPr = content_p.find(f"{{{W}}}pPr")
        if content_pPr is None:
            content_pPr = etree.SubElement(content_p, f"{{{W}}}pPr")
            content_p.insert(0, content_pPr)
        content_sectPr = etree.SubElement(content_pPr, f"{{{W}}}sectPr")
        pgSz = etree.SubElement(content_sectPr, f"{{{W}}}pgSz")
        pgSz.set(f"{{{W}}}w", "11906")

        # Now add an empty paragraph with its own sectPr
        empty_p = etree.SubElement(body, f"{{{W}}}p")
        empty_pPr = etree.SubElement(empty_p, f"{{{W}}}pPr")
        sectPr = etree.SubElement(empty_pPr, f"{{{W}}}sectPr")
        pgNumType = etree.SubElement(sectPr, f"{{{W}}}pgNumType")
        pgNumType.set(f"{{{W}}}fmt", "decimal")

        # Before fix: 2 paragraphs
        assert len(doc.paragraphs) == 2

        fix_section_breaks(doc)

        # After fix: the empty paragraph should be gone, only 1 paragraph remains
        assert len(doc.paragraphs) == 1

        # The content paragraph should now carry the merged sectPr
        content_p = body.find(f"{{{W}}}p")
        content_pPr = content_p.find(f"{{{W}}}pPr")
        assert content_pPr is not None
        moved_sectPr = content_pPr.find(f"{{{W}}}sectPr")
        assert moved_sectPr is not None
        moved_pgNumType = moved_sectPr.find(f"{{{W}}}pgNumType")
        assert moved_pgNumType is not None
        assert moved_pgNumType.get(f"{{{W}}}fmt") == "decimal"
