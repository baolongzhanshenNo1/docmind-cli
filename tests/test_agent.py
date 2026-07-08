"""Unit tests for pipeline/agent.py — DocMindAgent orchestration."""
import tempfile
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from pipeline.agent import DocMindAgent, AgentResult, _copy_format_params, _brief_action

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _make_minimal_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph('Hello')
    doc.save(str(path))


def _make_spec_docx(path: Path) -> None:
    """Create a minimal spec docx with H1 style defined."""
    from docx.shared import Pt
    doc = Document()
    doc.add_paragraph('Test')
    # Set some style properties
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    doc.save(str(path))


@pytest.fixture
def temp_target():
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        path = Path(f.name)
    _make_minimal_docx(path)
    yield path
    if path.exists():
        path.unlink()


@pytest.fixture
def temp_spec():
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        path = Path(f.name)
    _make_spec_docx(path)
    yield path
    if path.exists():
        path.unlink()


@pytest.fixture
def temp_output():
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        path = Path(f.name)
    yield path
    if path.exists():
        path.unlink()


# ── Test: AgentResult ───────────────────────────────────

class TestAgentResult:
    def test_default_values(self):
        r = AgentResult()
        assert r.plan == []
        assert r.issues == []
        assert r.fixed_count == 0
        assert r.remaining == 0
        assert r.output_path is None
        assert r.logs == []

    def test_with_values(self):
        from pipeline.fixer_v2 import FixerDiagnostic
        r = AgentResult(
            plan=[{'action': 'x'}],
            issues=[FixerDiagnostic(code='TEST', action_index=0, severity='info', detail='d')],
            fixed_count=5,
            remaining=1,
            output_path=Path('/tmp/test.docx'),
            logs=['log1', 'log2'],
        )
        assert r.fixed_count == 5
        assert r.remaining == 1
        assert r.output_path == Path('/tmp/test.docx')


# ── Test: DocMindAgent initialization ───────────────────

class TestDocMindAgentInit:
    def test_creates_instance(self):
        agent = DocMindAgent()
        assert agent is not None

    def test_initial_state(self):
        agent = DocMindAgent()
        assert agent._fix_plan == []
        assert agent._spec_rules == {}
        assert agent._user_overrides == {}

    def test_preview_plan_empty(self):
        agent = DocMindAgent()
        assert agent.preview_plan() == []

    def test_get_overrides_empty(self):
        agent = DocMindAgent()
        assert agent.get_overrides() == {}

    def test_reset_overrides(self):
        agent = DocMindAgent()
        agent._user_overrides['test'] = 'value'
        agent.reset_overrides()
        assert agent.get_overrides() == {}


# ── Test: DocMindAgent.run() ────────────────────────────

class TestDocMindAgentRun:
    def test_run_returns_agent_result(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        result = agent.run(
            spec_path=temp_spec,
            target_path=temp_target,
            output_path=temp_output,
        )
        assert isinstance(result, AgentResult)
        assert result.output_path == temp_output

    def test_run_produces_plan(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        result = agent.run(
            spec_path=temp_spec,
            target_path=temp_target,
            output_path=temp_output,
        )
        assert isinstance(result.plan, list)
        assert result.fixed_count >= 0

    def test_run_has_logs(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        result = agent.run(
            spec_path=temp_spec,
            target_path=temp_target,
            output_path=temp_output,
        )
        assert len(result.logs) > 0
        assert any('STEP 1/5' in l for l in result.logs)
        assert any('STEP 5/5' in l for l in result.logs)

    def test_run_output_exists(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(
            spec_path=temp_spec,
            target_path=temp_target,
            output_path=temp_output,
        )
        assert temp_output.exists()

    def test_run_with_same_output(self, temp_spec, temp_target):
        agent = DocMindAgent()
        result = agent.run(
            spec_path=temp_spec,
            target_path=temp_target,
            output_path=temp_target,  # same as target
        )
        assert result.output_path == temp_target


# ── Test: feedback loop ────────────────────────────────

class TestFeedbackLoop:
    def test_handle_feedback_key_value(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(spec_path=temp_spec, target_path=temp_target, output_path=temp_output)

        result = agent.handle_feedback('h1_font_east=黑体, h1_size_pt=16')
        assert isinstance(result, AgentResult)
        overrides = agent.get_overrides()
        assert overrides.get('h1_font_east') == '黑体'
        assert overrides.get('h1_size_pt') == 16

    def test_handle_feedback_numeric_conversion(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(spec_path=temp_spec, target_path=temp_target, output_path=temp_output)

        agent.handle_feedback('h1_size_pt=16.5')
        assert agent.get_overrides()['h1_size_pt'] == 16.5

    def test_handle_feedback_integer_conversion(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(spec_path=temp_spec, target_path=temp_target, output_path=temp_output)

        agent.handle_feedback('h1_size_pt=16')
        assert agent.get_overrides()['h1_size_pt'] == 16

    def test_handle_feedback_natural_language(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(spec_path=temp_spec, target_path=temp_target, output_path=temp_output)

        # Test blank page removal
        result = agent.handle_feedback('参考文献前面的空白页应该删除')
        assert isinstance(result, AgentResult)
        assert agent.get_overrides().get('remove_blank_pages') is True

    def test_handle_feedback_font_keyword(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(spec_path=temp_spec, target_path=temp_target, output_path=temp_output)

        agent.handle_feedback('把标题字体改成黑体')
        assert 'h1_font_east' in agent.get_overrides()
        assert agent.get_overrides()['h1_font_east'] == '黑体'

    def test_handle_feedback_multiple_rounds(self, temp_spec, temp_target, temp_output):
        agent = DocMindAgent()
        agent.run(spec_path=temp_spec, target_path=temp_target, output_path=temp_output)

        agent.handle_feedback('h1_font_east=黑体')
        assert agent.get_overrides()['h1_font_east'] == '黑体'

        agent.handle_feedback('body_font_east=宋体')
        assert agent.get_overrides()['body_font_east'] == '宋体'

        # Both overrides persist
        assert len(agent.get_overrides()) >= 2


# ── Test: helper functions ─────────────────────────────

class TestHelperFunctions:
    def test_copy_format_params(self):
        rules = {
            'font_east': '黑体',
            'font_ascii': 'Times New Roman',
            'font_size_pt': 16,
            'bold': True,
            'alignment': 'center',
        }
        params = {'style_id': 'test'}
        _copy_format_params(rules, params)
        assert params['font_eastAsia'] == '黑体'
        assert params['font_ascii'] == 'Times New Roman'
        assert params['font_size_pt'] == 16
        assert params['bold'] is True
        assert params['alignment'] == 'center'
        assert params['style_id'] == 'test'  # original key preserved

    def test_copy_format_params_partial(self):
        rules = {'font_east': '宋体'}
        params = {'style_id': 'test'}
        _copy_format_params(rules, params)
        assert params['font_eastAsia'] == '宋体'
        assert 'font_ascii' not in params

    def test_brief_action_set_style(self):
        result = _brief_action('set_style', {'style_id': '1', 'font_ascii': 'Arial'})
        assert '1' in result
        assert 'font_ascii=Arial' in result

    def test_brief_action_set_sectpr_type(self):
        result = _brief_action('set_sectpr_type', {'section_index': 0, 'val': 'nextPage'})
        assert 'sec[0]' in result
        assert 'nextPage' in result

    def test_brief_action_set_header_font(self):
        result = _brief_action('set_header_font', {
            'header_path': 'word/header2.xml',
            'font_eastAsia': '宋体',
            'font_ascii': 'Arial',
        })
        assert 'word/header2.xml' in result
        assert 'font=宋体/Arial' in result

    def test_brief_action_add_page_number(self):
        result = _brief_action('add_page_number', {
            'footer_path': 'word/footer2.xml',
            'format': '第{PAGE}页',
        })
        assert 'word/footer2.xml' in result
        assert '第{PAGE}页' in result

    def test_brief_action_remove_extra_sectpr(self):
        result = _brief_action('remove_extra_sectpr', {})
        assert 'cleanup' in result


# ── Test: reconciliation produces valid plan ────────────

class TestReconciliation:
    def test_reconcile_produces_non_empty_plan(self, temp_spec, temp_target):
        agent = DocMindAgent()
        agent._spec_path = temp_spec
        agent._target_path = temp_target
        agent._spec_rules = {
            'h1': {'font_east': '黑体', 'font_size_pt': 16, 'bold': True},
            'body': {'font_east': '宋体', 'font_ascii': 'Times New Roman'},
        }
        agent._style_ids = {'h1': '1', 'body': 'a8'}
        agent._heading_info = {'h1_count': 1, 'h2_count': 0, 'h3_count': 0}
        agent._section_count = 1

        plan = agent._reconcile()
        assert len(plan) > 0

        # Should include at least style changes
        style_actions = [a for a in plan if a['action'] == 'set_style']
        assert len(style_actions) >= 1

    def test_reconcile_no_spec_rules(self, temp_spec, temp_target):
        agent = DocMindAgent()
        agent._spec_path = temp_spec
        agent._target_path = temp_target
        agent._spec_rules = {}
        agent._style_ids = {}
        agent._heading_info = {}
        agent._section_count = 1

        plan = agent._reconcile()
        # No style rules → no style actions, but still may have other actions
        style_actions = [a for a in plan if a['action'] == 'set_style']
        assert len(style_actions) == 0


# ── Test: user overrides applied to plan ────────────────

class TestUserOverrides:
    def test_override_h1_font_applied(self, temp_spec, temp_target):
        agent = DocMindAgent()
        agent._spec_path = temp_spec
        agent._target_path = temp_target
        agent._spec_rules = {
            'h1': {'font_east': '宋体', 'font_size_pt': 14},
            'body': {},
        }
        agent._style_ids = {'h1': '1', 'body': 'a8'}
        agent._heading_info = {'h1_count': 1}
        agent._section_count = 1
        agent._user_overrides = {'h1_font_east': '黑体'}

        plan = agent._reconcile()

        # Find h1 style action
        h1_actions = [a for a in plan if a['action'] == 'set_style' and a['params'].get('style_id') == '1']
        assert len(h1_actions) >= 1
        # The override should have changed the font
        assert h1_actions[0]['params'].get('font_eastAsia') == '黑体'

    def test_override_header_font_applied(self, temp_spec, temp_target):
        agent = DocMindAgent()
        agent._spec_path = temp_spec
        agent._target_path = temp_target
        agent._spec_rules = {
            'body': {'font_east': '宋体', 'font_ascii': 'Arial'},
        }
        agent._style_ids = {'body': 'a8'}
        agent._heading_info = {'h1_count': 0}
        agent._section_count = 1
        agent._user_overrides = {
            'header_font_east': '黑体',
            'header_font_ascii': 'Times New Roman',
        }

        plan = agent._reconcile()

        # Find header font actions
        hdr_actions = [a for a in plan if a['action'] == 'set_header_font']
        for ha in hdr_actions:
            assert ha['params'].get('font_eastAsia') == '黑体'
            assert ha['params'].get('font_ascii') == 'Times New Roman'
