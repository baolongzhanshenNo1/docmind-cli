"""python-docx post-processing: rebuild all section headers and footers."""
from pathlib import Path


def rebuild_headers_footers(docx_path: Path, section_headers: list[dict]) -> None:
    """用 python-docx 重建所有节的页眉（含下划线）+ PAGE 域页脚。

    section_headers: [{"section_index": N, "default_text": "...", "even_text": None}, ...]
    """
    from docx import Document
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(docx_path))

    for sh in section_headers:
        si = sh.get("section_index", sh.get("sectpr_index", -1))
        text = sh.get("default_text", "")
        if not text or si < 0 or si >= len(doc.sections):
            continue

        sec = doc.sections[si]

        # ── 断链接 ──
        try:
            sec.header.is_linked_to_previous = False
            sec.footer.is_linked_to_previous = False
        except Exception:
            pass

        # ── 页眉：居中 + 下划线 + 宋体/Times New Roman ──
        hdr = sec.header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        # East-Asian font
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), '宋体')

        # Bottom border (underline)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)

        # ── 页脚：居中 PAGE 域 ──
        ftr = sec.footer
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()

        run_prefix = fp.add_run('第')
        run_prefix.font.name = 'Times New Roman'
        rPr_p = run_prefix._r.get_or_add_rPr()
        rFs = rPr_p.find(qn('w:rFonts'))
        if rFs is None:
            rFs = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
            rPr_p.insert(0, rFs)
        else:
            rFs.set(qn('w:eastAsia'), '宋体')

        # PAGE field
        fld_begin = parse_xml(
            f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>'
        )
        fld_instr = parse_xml(
            f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        )
        fld_sep = parse_xml(
            f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="separate"/></w:r>'
        )
        fld_val = parse_xml(
            f'<w:r {nsdecls("w")}><w:t>1</w:t></w:r>'
        )
        fld_end = parse_xml(
            f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>'
        )
        fp._p.append(fld_begin)
        fp._p.append(fld_instr)
        fp._p.append(fld_sep)
        fp._p.append(fld_val)
        fp._p.append(fld_end)

        run_suffix = fp.add_run('页')
        run_suffix.font.name = 'Times New Roman'
        rPr_s = run_suffix._r.get_or_add_rPr()
        rFs2 = rPr_s.find(qn('w:rFonts'))
        if rFs2 is None:
            rFs2 = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
            rPr_s.insert(0, rFs2)
        else:
            rFs2.set(qn('w:eastAsia'), '宋体')

    doc.save(str(docx_path))
