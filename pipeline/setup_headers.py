"""
流水线步骤2: 页眉页脚 — 页眉文字 + 页码 + 页边距

合并原来 fix_v10/v11/fix_headers 的功能。
"""
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from .config import PipelineConfig

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def _set_header(sec, text: str, font_size_pt: float = 9):
    """设置一个 section 的默认页眉"""
    h = sec.header
    h.is_linked_to_previous = False
    for p in list(h.paragraphs):
        p._element.getparent().remove(p._element)
    if not h.paragraphs:
        h.add_paragraph()
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(font_size_pt)
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(run._element, qn('w:rPr'))
    # 下划线
    if rPr.find(qn('w:pBdr')) is None and rPr.find(qn('w:bdr')) is None:
        pass  # border 由其他逻辑处理
    return h


def run(config: PipelineConfig):
    """设置所有节的页眉、页脚、页码"""
    print('[setup_headers] Loading...')
    doc = docx.Document(str(config.input_docx))
    
    # ── 页眉 ──
    print('[setup_headers] Setting headers...')
    for i, sec in enumerate(doc.sections):
        hdr_info = config.section_headers.get(i)
        if hdr_info is None:
            continue
        
        default_text, even_text, first_text = hdr_info
        header_font_size = 9  # 统一字号
        
        # 断开链接
        sec.header.is_linked_to_previous = False
        
        # 设置默认页眉
        if default_text:
            _set_header(sec, default_text, header_font_size)
        
        # 设置偶数页页眉（如果需要）
        if even_text:
            try:
                even_hdr = sec.even_page_header
                even_hdr.is_linked_to_previous = False
                for p in list(even_hdr.paragraphs):
                    p._element.getparent().remove(p._element)
                if not even_hdr.paragraphs:
                    even_hdr.add_paragraph()
                p = even_hdr.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(even_text)
                run.font.size = Pt(header_font_size)
            except Exception as e:
                print(f'  WARNING: Sec {i} even header failed: {e}')
        
        # 设置首页页眉（如果需要）
        if first_text:
            try:
                first_hdr = sec.first_page_header
                first_hdr.is_linked_to_previous = False
                for p in list(first_hdr.paragraphs):
                    p._element.getparent().remove(p._element)
                if not first_hdr.paragraphs:
                    first_hdr.add_paragraph()
                p = first_hdr.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(first_text)
                run.font.size = Pt(header_font_size)
            except Exception as e:
                print(f'  WARNING: Sec {i} first header failed: {e}')
    
    # ── 页脚（页码） ──
    print('[setup_headers] Setting footers...')
    footer_font_size = config.footer_font_size_pt
    for i, sec in enumerate(doc.sections):
        if i == 0:
            continue  # 封面无页码
        
        for footer_attr in ['footer', 'even_page_footer']:
            try:
                ftr = getattr(sec, footer_attr) if footer_attr == 'footer' else sec.even_page_footer
            except Exception:
                continue
            
            ftr.is_linked_to_previous = False
            for p in list(ftr.paragraphs):
                p._element.getparent().remove(p._element)
            
            p = ftr.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # '第'
            run1 = p.add_run('第')
            run1.font.size = Pt(footer_font_size)
            
            # PAGE field
            run2 = p.add_run()
            run2.font.size = Pt(footer_font_size)
            etree.SubElement(run2._element, qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
            
            run3 = p.add_run()
            run3.font.size = Pt(footer_font_size)
            etree.SubElement(run3._element, qn('w:instrText'),
                           {'{http://www.w3.org/XML/1998/namespace}space': 'preserve'}).text = ' PAGE '
            
            run4 = p.add_run()
            run4.font.size = Pt(footer_font_size)
            etree.SubElement(run4._element, qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
            
            run5 = p.add_run('1')
            run5.font.size = Pt(footer_font_size)
            
            run6 = p.add_run()
            run6.font.size = Pt(footer_font_size)
            etree.SubElement(run6._element, qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
            
            # '页'
            run7 = p.add_run('页')
            run7.font.size = Pt(footer_font_size)
    
    # ── 页码类型（罗马数字前部 + 阿拉伯正文） ──
    for i, sec in enumerate(doc.sections):
        sp = sec._sectPr
        if i == 0:
            continue  # 封面无页码
        elif i <= 3:  # 摘要/ABSTRACT/目录 → 罗马
            fmt = 'upperRoman'
        else:
            fmt = 'decimal'
        
        old_pg = sp.find(qn('w:pgNumType'))
        if old_pg is not None:
            sp.remove(old_pg)
        
        attrs = {qn('w:fmt'): fmt}
        if i == 1 or i == 4:  # 起始页
            attrs[qn('w:start')] = '1'
        etree.SubElement(sp, qn('w:pgNumType'), attrs)
    
    # ── 页边距 ──
    for sec in doc.sections:
        sp = sec._sectPr
        pm = sp.find(qn('w:pgMar'))
        if pm is not None:
            pm.set(qn('w:header'), str(config.page.header_distance_twips))
            pm.set(qn('w:footer'), str(config.page.footer_distance_twips))
    
    doc.save(str(config.input_docx))
    print(f'[setup_headers] Saved to {config.input_docx}')
